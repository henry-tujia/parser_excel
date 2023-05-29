from docx import Document
from docx.shared import Inches
from parseMain import parse_main
import os

def exportword(filename,year,month,day,root_path,combine_name,dubget_name):
    if "\\" in root_path:
        root_path = root_path.replace("\\","/")

    document = Document()
    from docx.oxml.ns import qn

    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    paras,sheet = parse_main(root_path,combine_name,dubget_name,year)

    columns = ['项目',	 '本年预算' ,	'截至本期实际发生',	'执行率（%）']

    assert sheet.shape[0] != len(columns)

    document.add_heading('山东蓝色经济创业投资有限公司\n{0}年{1}月份预算执行分析报告'.format(year,month), 0)

    document.add_heading('一、	经济效益完成情况', 1)

    for heading,para in paras.items():
        document.add_heading(heading, 2)
        document.add_paragraph(para)
        # document.add_paragraph(para.replace('2022年', f"{str(year)}年"))
        if heading=="管理费用":
            document.add_paragraph('本期费用明细如下：')
            table = document.add_table(1,sheet.shape[1],style = 'Light Shading Accent 1')
            for i in range(sheet.shape[1]):
                hdr_cells = table.rows[0].cells
                hdr_cells[i].text = columns[i]
            for index,row in sheet.iterrows():
                row_cells = table.add_row().cells
                values = row.values
                for i in range(sheet.shape[1]):
                    value = values[i]
                    if i == 2 and float(value) == 0:
                        row_cells[i].text = " - "
                        continue
                    if isinstance(value,str):
                        row_cells[i].text = values[i]
                    else:
                        row_cells[i].text = format(values[i],'.2f')
    document.add_heading('二、	工作建议',1)

    document.add_paragraph('{0}年，蓝色创投各部门厉行节支增效原则，开源节流，降低费用支出；结合自身实际，围绕战略方向和经营管理目标，梳理业务流程，加强风险管理和制度体系建设，不断夯实基础管理工作；抓住转型机遇，积极拓展涉海涉蓝市场业务，加大项目储备力度。'.format(year))

    p = document.add_paragraph('山东蓝色经济创业投资有限公司')
    p.alignment = 2

    if int(month)==12:
        year = int(year)+1

    p = document.add_paragraph('{0}年{1}月{2}日'.format(year,int(month)%12+1,day))
    p.alignment = 2

    document.add_page_break()

    document.save(os.path.join(root_path,filename).replace('\\', '/'))